import docx, time
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt

tag = True

while tag:
    print('欢迎进入迪庆下料系统')
    type = input('请选择下料产品：(1:麻花钢，2:插接，q：退出)').strip()
    if type == '1':
        print('您选择的是普通麻花钢，请输入相关参数（单位：mm）')

        while tag:
            kjzxj = input('请输入孔距中心距>>:').strip()

            if not kjzxj.isdigit():
                print('输入错误，请输入数字')
                continue
            else:
                kjzxj = int(kjzxj)
                break

        while tag:
            bbzrsc = input('请输入包边自然缩尺>>:').strip()

            if not bbzrsc.isdigit():
                print('输入错误，请输入数字')
                continue
            else:
                bbzrsc = int(bbzrsc)
                break

        while tag:
            bgfxzrsc = input('请输入扁钢方向自然缩尺>>:').strip()

            if not bgfxzrsc.isdigit():
                print('输入错误，请输入数字')
                continue
            else:
                bgfxzrsc = int(bgfxzrsc)
                break

        while tag:
            bgfx = input('请输入扁钢方向>>:').strip()

            if not bgfx.isdigit():
                print('输入错误，请输入数字')
                continue
            else:
                bgfx = int(bgfx)
                break
        while tag:
            mhgfx = input('请输入麻花钢方向>>:').strip()

            if not mhgfx.isdigit():
                print('输入错误，请输入数字')
                continue
            else:
                mhgfx = int(mhgfx)
                break
        while tag:
            ks = input('请输入块数>>:').strip()

            if not ks.isdigit():
                print('输入错误，请输入数字')
                continue
            else:
                ks = int(ks)
                break

        print('正在计算，请稍后')
        res = {
            '扁钢方向尺寸': bgfx - bgfxzrsc,
            '扁钢条根数': (mhgfx // kjzxj + 1) * ks,
            '包边尺寸': mhgfx - bbzrsc,
            '边条根数': ks * 2
        }

        print(res)
        name = input('请输入客户公司名称').strip()

        start = time.time()
        print('正在写入文件，请稍后')
        # 创建内存中的word文档对象
        file = docx.Document()
        # 写入若干段落
        # title = file.add_heading("河北逍迪丝网制品有限公司", 2)
        # title = title.add_run("河北逍迪丝网制品有限公司")
        # title.font.bold = True

        # style = file.styles.add_style('Indent', WD_STYLE_TYPE.PARAGRAPH)
        # paragraph_format = style.paragraph_format
        # paragraph_format.left_indent = Inches(0.25)
        # paragraph_format.first_line_indent = Inches(-0.25)
        # paragraph_format.space_before = Pt(12)
        # paragraph_format.widow_control = True
        # print(paragraph_format)

        file.add_heading("                                   河北逍迪丝网制品有限公司", level=1)

        file.add_heading("                                                                                                                    下料单", level=5)
        file.add_heading("材质：                                                                           工期：", level=5)
        file.add_heading("规格：", level=5)
        file.add_heading("孔距中心距：%s                                                                   麻花钢：" % kjzxj, level=5)
        file.add_heading("扁钢厚度：                                                                         圆钢：", level=5)
        file.add_heading("包边自然缩尺：%s                                                               齿形：" % bbzrsc, level=5)
        file.add_heading("扁钢方向自然缩尺：%s" % bgfxzrsc, level=5)
        file.add_heading("扁钢宽度：", level=5)
        file.add_heading("表面处理：", level=5)
        file.add_heading("——————————————————————————————————————————", level=5)
        file.add_heading("公司名称：%s" % name, level=5)
        file.add_heading("联系电话：", level=5)
        file.add_heading("收货地址：", level=5)
        file.add_heading("运费哪方出：", level=5)
        file.add_heading("收货人：", level=5)
        file.add_heading("收货人电话：", level=5)
        file.add_heading("货款定金：", level=5)
        file.add_heading("含税总价格：                元                                                  不含税总价格：         元", level=5)
        file.add_heading("收款方式：", level=5)
        file.add_heading("交易方式：", level=5)
        file.add_heading("冲豆裁料工费：             平米        元       踏步        元       沟盖板       元", level=5)
        file.add_heading("焊料费费：                    平米        元       踏步        元       沟盖板       元", level=5)
        file.add_heading("焊工工费：                 平米        元       踏步        元       沟盖板       元 ", level=5)
        file.add_heading("业务员：                     提成：        元", level=5)
        file.add_heading("——————————————————————————————————————————", level=5)
        file.add_heading("附属：", level=5)
        file.add_heading("——————————————————————————————————————————", level=5)
        table = file.add_table(rows=1, cols=8, style='Table Grid')  # 创建带边框的表格
        hdr_cells = table.rows[0].cells  # 获取第0行所有所有单元格
        hdr_cells[0].text = '外形尺寸扁钢方向'
        hdr_cells[1].text = '外形尺寸麻花钢方向'
        hdr_cells[2].text = '块数'
        hdr_cells[3].text = '扁钢方向尺寸（含两边）'
        hdr_cells[4].text = '扁钢条根数'
        hdr_cells[5].text = '包边尺寸'
        hdr_cells[6].text = '边条根数'
        hdr_cells[7].text = '附属'

        cells = table.add_row().cells
        cells[0].text = '%s' % bgfx
        cells[1].text = '%s' % mhgfx
        cells[2].text = '%s' % ks
        cells[3].text = '%s' % res['扁钢方向尺寸']
        cells[4].text = '%s' % res['扁钢条根数']
        cells[5].text = '%s' % res['包边尺寸']
        cells[6].text = '%s' % res['边条根数']
        cells[7].text = ' '

        file.add_heading("合计：%s块" % ks)
        now = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time()))
        file.add_heading("出单时间：%s" % now)

        # 保存
        file.save("%s.docx" % name)
        end = time.time()
        print('文件保存完成，用时%s秒,请在左侧文件目录或者桌面快捷方式查看' % (end - start))




    elif type == '2':
        print('您选择的是插接类型，请输入相关参数（单位：mm）')

        while tag:
            bgzxj = input('请输入扁钢中心距>>:').strip()

            if not bgzxj.isdigit():
                print('输入错误，请输入数字')
                continue
            else:
                bgzxj = int(bgzxj)
                break

        while tag:
            bbzxj = input('请输入包边中心距>>:').strip()

            if not bbzxj.isdigit():
                print('输入错误，请输入数字')
                continue
            else:
                bbzxj = int(bbzxj)
                break

        while tag:
            bbzrsc = input('请输入包边自然缩尺>>:').strip()

            if not bbzrsc.isdigit():
                print('输入错误，请输入数字')
                continue
            else:
                bbzrsc = int(bbzrsc)
                break

        while tag:
            bgfxzrsc = input('请输入扁钢方向自然缩尺>>:').strip()

            if not bgfxzrsc.isdigit():
                print('输入错误，请输入数字')
                continue
            else:
                bgfxzrsc = int(bgfxzrsc)
                break

        while tag:
            bgfx = input('请输入扁钢方向>>:').strip()

            if not bgfx.isdigit():
                print('输入错误，请输入数字')
                continue
            else:
                bgfx = int(bgfx)
                break

        while tag:
            bbfx = input('请输入包边方向>>:').strip()
            if not bbfx.isdigit():
                print('输入错误，请输入数字')
                continue
            else:
                bbfx = int(bbfx)
                break

        while tag:
            ks = input('请输入块数>>:').strip()
            if not ks.isdigit():
                print('输入错误，请输入数字')
                continue
            else:
                ks = int(ks)
                break

        print('正在计算，请稍后')
        res = {
            '尺寸(含边)': bgfx - bgfxzrsc,
            '扁钢条根数': (bbfx // bgzxj + 1) * ks,
            '尺寸': bbfx - bgfxzrsc,
            '内条根数': (bgfx // bbzxj - 1) * ks,
            '包边': bbfx - bbzrsc,
            '边条根数': ks * 2
        }

        print(res)

        name = input('请输入客户公司名称').strip()

        start = time.time()
        print('正在写入文件，请稍后')
        # 创建内存中的word文档对象
        file = docx.Document()
        # 写入若干段落
        file.add_heading("                                   河北逍迪丝网制品有限公司", level=1)
        file.add_heading("                                                                                                                    下料单", level=5)
        file.add_heading("材质：                                                                           工期：", level=5)
        file.add_heading("规格：                                                                                   插接：是", level=5)
        file.add_heading("扁钢中心距：%s                                                                   齿形：" % bgzxj, level=5)
        file.add_heading("包边中心距：%s" % bbzxj, level=5)
        file.add_heading("包边自然缩尺：%s" % bbzrsc, level=5)
        file.add_heading("扁钢方向自然缩尺：%s" % bgfxzrsc, level=5)
        file.add_heading("扁钢厚度：", level=5)
        file.add_heading("扁钢宽度：", level=5)
        file.add_heading("表面处理：", level=5)
        file.add_heading("——————————————————————————————————————————", level=5)
        file.add_heading("公司名称：%s" % name, level=5)
        file.add_heading("联系电话：", level=5)
        file.add_heading("收货地址：", level=5)
        file.add_heading("运费哪方出：", level=5)
        file.add_heading("收货人：", level=5)
        file.add_heading("收货人电话：", level=5)
        file.add_heading("货款定金：", level=5)
        file.add_heading("含税总价格：            元                                      不含税总价格：           元", level=5)
        file.add_heading("收款方式：", level=5)
        file.add_heading("交易方式：", level=5)
        file.add_heading("冲豆裁料工费：                  平米         元           踏步        元       沟盖板        元", level=5)
        file.add_heading("焊工工费：                   平米          元           踏步         元        沟盖板          元 ", level=5)
        file.add_heading("焊材费：                      平米          元          踏步         元         沟盖板       元 ", level=5)
        file.add_heading("业务员：                      提成：            元", level=5)
        file.add_heading("——————————————————————————————————————————", level=5)
        file.add_heading("附属：", level=5)
        file.add_heading("——————————————————————————————————————————", level=5)
        table = file.add_table(rows=1, cols=10, style='Table Grid')  # 创建带边框的表格
        hdr_cells = table.rows[0].cells  # 获取第0行所有所有单元格
        hdr_cells[0].text = '外形尺寸扁钢方向'
        hdr_cells[1].text = '外形尺寸包边方向'
        hdr_cells[2].text = '块数'
        hdr_cells[3].text = '尺寸（含边）'
        hdr_cells[4].text = '扁钢条根数'
        hdr_cells[5].text = '尺寸'
        hdr_cells[6].text = '内条根数'
        hdr_cells[7].text = '包边'
        hdr_cells[8].text = '边条根数'
        hdr_cells[9].text = '附属'

        cells = table.add_row().cells
        cells[0].text = '%s' % bgfx
        cells[1].text = '%s' % bbfx
        cells[2].text = '%s' % ks
        cells[3].text = '%s' % res['尺寸(含边)']
        cells[4].text = '%s' % res['扁钢条根数']
        cells[5].text = '%s' % res['尺寸']
        cells[6].text = '%s' % res['内条根数']
        cells[7].text = '%s' % res['包边']
        cells[8].text = '%s' % res['边条根数']
        cells[9].text = ' '

        file.add_heading("合计：%s块" % ks)
        now = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time()))
        file.add_heading("出单日期：%s" % now)

        # 保存
        file.save("%s.docx" % name)
        end = time.time()
        print('文件保存完成，用时%s秒,请在左侧文件目录或者桌面快捷方式查看' % (end - start))
    elif type == 'q':
        break
    else:
        print('输入1或者2选择，您的输入有误，请重新输入')
